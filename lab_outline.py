import docx
import subprocess


def getDeliverables(path):
    deliverables = []
    with open(path, "r") as f:
        for line in f.readlines():
            if "This criterion is linked to a Learning Outcome" in line:
                deliverable = line.split("This criterion is linked to a Learning Outcome")[1] # the first index is blank second is the deliverable
                deliverables.append(deliverable)

    
    return deliverables


def makeLabOutline(deliverables):

    outline = docx.Document() # create outline document
    outline.add_paragraph("Galen Dively") # add my name

    table = outline.add_table(len(deliverables)+1, 2) # create the table with the length of x deliverables + 1 for titles

    # add title of columns
    table.rows[0].cells[0].text = "Requirements"
    table.rows[0].cells[1].text = "Submission"

    count = 0 # keeps count of current deliverable
    # loop through each row in the table
    for row in table.rows:
        # skip the first cell
        if row.cells[0].text == "Requirements":
            continue

        row.cells[0].text = deliverables[count] # add the deliverble from the deliverables array
        count +=1 # increase the deliverable count by 1

    title = input("Please Enter Title: ")
    outline.save(f"output/{title}.docx") # save it to users Homework folder
    return f"output/{title}.docx"


# libreoffice must be installed and located in path

def openDocument(output_path):
    run_command = f"libreoffice"
    subprocess.run([run_command, output_path])


if __name__ == "__main__":
    # lab: the orginal lab
    # outline: lab with table and questions
    print("Starting")
    open_after_creation = input("Do you want to open in libreoffice?(Y/n)").upper()
    match open_after_creation:
        case "Y":
            open_after_creation = True
        case "N":
            open_after_creation = False

    path = "input/assignment.txt"
    deliverables = getDeliverables(path)
    output_path = makeLabOutline(deliverables)
    if open_after_creation:
        print("Opening")
        openDocument(output_path)
    print("Complete")
    quit()
